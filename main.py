import PyPDF2
from docx import Document
from googletrans import Translator

# Создаем объект класса Translator
translator = Translator(service_urls=['translate.google.com'])

# Открыть PDF-файл
pdf_file = open('book.pdf', 'rb')

# Создать объект PdfReader
pdf_reader = PyPDF2.PdfReader(pdf_file)

# Создать объект Document для записи текста в документ Word
docx_document_to_word = Document()

# Обойти все страницы PDF-файла и извлечь текст
for page_num in range(len(pdf_reader.pages)):
    try:
        page = pdf_reader.pages[page_num]
        text = page.extract_text()
    except:
        print(f'Ошибка при обработке страницы {page_num+1}, страница пропущена')
        continue

    # Добавить текст в документ Word
    if page_num == 0:
        docx_document_to_word.add_heading('Document Title', 0)
    else:
        docx_document_to_word.add_page_break()
    docx_document_to_word.add_paragraph(f'Page {page_num + 1}:')
    try:
        text = translator.translate(text, dest='ru')
        docx_document_to_word.add_paragraph(text.text)
    except Exception as e:
        print(f'Ошибка при сохранении страницы {page_num + 1}, страница пропущена. Ошибка {e}')
        docx_document_to_word.add_page_break()
        continue

    print(f'Страница {page_num+1} сохранена в формате docx.')

# Сохранить документ Word
docx_document_to_word.save('file.docx')

# Закрыть файлы
pdf_file.close()